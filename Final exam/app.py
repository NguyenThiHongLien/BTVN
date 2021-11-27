from flask import Flask, render_template, request, redirect, url_for, flash
from flask_mysqldb import MySQL
from os import error
from docx import Document


app = Flask(__name__)
app.secret_key = 'many random bytes'


app.config['MYSQL_HOST'] = 'remotemysql.com'
app.config['MYSQL_USER'] = 'QyYiCxuWS2'
app.config['MYSQL_PASSWORD'] = 'DIeOFa6xzf'
app.config['MYSQL_DB'] = 'QyYiCxuWS2'

mysql = MySQL(app)

@app.route('/')
def Index():
    cur = mysql.connection.cursor()
    cur.execute("SELECT  * FROM blogs Order by created_time desc")
    blogs = cur.fetchall()
    cur.close()
    


    return render_template('index.html', blogs=blogs )

@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/resignation")
def resignation():
    return render_template("resignation.html")

@app.route('/add', methods = ['POST'])
def insert():

    if request.method == "POST":
        flash("Data Inserted Successfully")
        title = request.form['title']
        content = request.form['content']
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO blogs (title, content) VALUES (%s, %s)", (title, content))
        mysql.connection.commit()
        return redirect(url_for('Index'))




@app.route("/<id>", methods = ['GET'])
def edit(id):
    cur = mysql.connection.cursor()
    cur.execute("SELECT  * FROM blogs WHERE id = %s", [id])
    post = cur.fetchone()
    cur.close()
    
    return render_template("edit.html", post =post)




@app.route("/post/<post_id>", methods=["POST"])
def edit_post(post_id):
    title = request.form.get("title")
    content = request.form.get("content")

    if title and content:
        sql = '''
            SELECT * FROM blogs
            WHERE id = %s
        '''
        cur = mysql.connection.cursor()
        cur.execute(sql, [post_id])
        post = cur.fetchone()

        if not post:
            flash("Post does not exists")
        else:
            sql = '''
                UPDATE blogs
                SET title = %s,
                    content = %s
                WHERE
                    id = %s
            '''

            cur.execute(sql, (title, content, post_id))
            mysql.connection.commit()
            flash("Post updated!")

        return redirect(url_for('Index'))
    else:
        flash("Title and content cannot be empty")
        return redirect(url_for('Index'))

@app.route("/resignation-letter", methods=["GET"])
def render_form():
    return render_template("resignation.html")



@app.route("/resignation-letter", methods=["POST"])
def create_file():
    from pathlib import Path
    from docx import Document
    from docx.shared import Cm
    fullname = request.form.get("fullname")
    reason = request.form.get("reason")

    if fullname and reason:
        doc = Document()
        doc.add_heading("Resignation Letter", level=0)
        doc.add_paragraph(fullname)
        doc.add_paragraph(reason)
        file_name = 'resignation-letter.docx'
        doc.save(str(Path(__file__).parent.absolute())+'/static/resignation-letter.docx')
        return render_template("resignation.html", file_created=True)
    else:
        flash("Fullname and reason cannot be empty")
        return render_template("resignation.html", file_created=False)

if __name__ == "__main__":
    app.run(debug=True)